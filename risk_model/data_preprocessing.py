"""
Data preprocessing module for DPR documents
"""
import os
import re
import string
from typing import Dict, List, Any, Tuple, Optional, Union

import numpy as np
import pandas as pd
import nltk
import spacy
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import word_tokenize
import PyPDF2
import docx
import openpyxl

# Download NLTK resources
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
    nltk.data.find('corpora/wordnet')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('wordnet')

# Initialize the SpaCy model (you'll need to download this with: python -m spacy download en_core_web_sm)
try:
    nlp = spacy.load('en_core_web_sm')
except OSError:
    import sys
    print("SpaCy model 'en_core_web_sm' is not installed. Please install it using:")
    print("python -m spacy download en_core_web_sm")
    sys.exit(1)


class DPRDocumentProcessor:
    """
    A class for processing DPR documents in various formats and extracting
    structured information for risk analysis.
    """
    
    def __init__(self, config: Dict[str, Any]):
        """
        Initialize the document processor.
        
        Args:
            config: Configuration dictionary with processing parameters
        """
        self.config = config
        self.stopwords = set(stopwords.words('english'))
        self.lemmatizer = WordNetLemmatizer()
    
    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text content from a file based on its format.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Extracted text content as a string
            
        Raises:
            ValueError: If the file format is not supported
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in self.config["SUPPORTED_FILE_TYPES"]:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Check file size
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > self.config["MAX_FILE_SIZE_MB"]:
            raise ValueError(f"File size ({file_size_mb:.2f} MB) exceeds the maximum allowed size ({self.config['MAX_FILE_SIZE_MB']} MB)")
        
        # Extract text based on file type
        if file_ext == ".pdf":
            return self._extract_text_from_pdf(file_path)
        elif file_ext == ".docx":
            return self._extract_text_from_docx(file_path)
        elif file_ext == ".xlsx":
            return self._extract_text_from_xlsx(file_path)
        elif file_ext == ".txt":
            return self._extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + " "
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            raise
        return text
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + " "
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            raise
        return text
    
    def _extract_text_from_xlsx(self, file_path: str) -> str:
        """Extract text from XLSX files"""
        text = ""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value:
                            text += str(cell.value) + " "
        except Exception as e:
            print(f"Error extracting text from XLSX: {e}")
            raise
        return text
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with another encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                print(f"Error extracting text from TXT: {e}")
                raise
    
    def preprocess_text(self, text: str) -> str:
        """
        Preprocess the extracted text.
        
        Args:
            text: Raw text from a document
            
        Returns:
            Preprocessed text
        """
        if not text:
            return ""
        
        # Convert to lowercase
        text = text.lower()
        
        # Remove special characters and digits
        text = re.sub(r'[^\w\s]', ' ', text)
        text = re.sub(r'\d+', ' ', text)
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def tokenize_and_lemmatize(self, text: str) -> List[str]:
        """
        Tokenize and lemmatize text.
        
        Args:
            text: Preprocessed text
            
        Returns:
            List of tokens
        """
        # Tokenize
        tokens = word_tokenize(text)
        
        # Remove stopwords and lemmatize
        tokens = [
            self.lemmatizer.lemmatize(token) 
            for token in tokens 
            if token not in self.stopwords and len(token) > 2
        ]
        
        return tokens
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """
        Extract different sections from the document.
        
        Args:
            text: Preprocessed document text
            
        Returns:
            Dictionary with section names as keys and section content as values
        """
        # Common sections in DPR documents
        section_patterns = {
            "executive_summary": r'executive\s+summary|summary',
            "project_description": r'project\s+description|about\s+the\s+project',
            "budget_estimation": r'budget|cost\s+estimation|financial\s+plan',
            "technical_specifications": r'technical\s+specifications|technical\s+details',
            "implementation_plan": r'implementation\s+plan|execution\s+plan',
            "environmental_impact": r'environmental\s+impact|environmental\s+assessment',
            "risk_assessment": r'risk\s+assessment|risk\s+analysis'
        }
        
        sections = {}
        doc = nlp(text)
        
        # Process each section using SpaCy for better paragraph detection
        for section_name, pattern in section_patterns.items():
            # Find all matches for the section header
            matches = list(re.finditer(pattern, text, re.IGNORECASE))
            
            if matches:
                # Take the first match for simplicity
                start_idx = matches[0].end()
                
                # Try to find the next section start
                next_section_starts = []
                for next_pattern in section_patterns.values():
                    if next_pattern != pattern:
                        next_matches = list(re.finditer(next_pattern, text[start_idx:], re.IGNORECASE))
                        if next_matches:
                            next_section_starts.append(start_idx + next_matches[0].start())
                
                # Get the content between this section start and the next section start
                if next_section_starts:
                    end_idx = min(next_section_starts)
                    sections[section_name] = text[start_idx:end_idx].strip()
                else:
                    sections[section_name] = text[start_idx:].strip()
        
        return sections
    
    def extract_risk_indicators(self, text: str, risk_keywords: Dict[str, List[str]]) -> Dict[str, float]:
        """
        Extract risk indicators from the text based on keywords.
        
        Args:
            text: Preprocessed document text
            risk_keywords: Dictionary of risk categories and their associated keywords
            
        Returns:
            Dictionary with risk categories and their normalized frequency scores
        """
        tokens = self.tokenize_and_lemmatize(text)
        token_count = len(tokens)
        
        if token_count == 0:
            return {category: 0.0 for category in risk_keywords}
        
        risk_indicators = {}
        
        for category, keywords in risk_keywords.items():
            # Count occurrences of each keyword
            keyword_count = sum(tokens.count(keyword) for keyword in keywords)
            
            # Normalize by token count and keyword list length
            normalized_score = keyword_count / (token_count * len(keywords))
            
            # Apply sigmoid function to smooth scores between 0 and 1
            risk_indicators[category] = 1 / (1 + np.exp(-10 * (normalized_score - 0.01)))
        
        return risk_indicators
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract structured information.
        
        Args:
            file_path: Path to the DPR document
            
        Returns:
            Dictionary with processed document information
        """
        # Extract and preprocess text
        raw_text = self.extract_text_from_file(file_path)
        preprocessed_text = self.preprocess_text(raw_text)
        
        # Extract sections
        sections = self.extract_sections(preprocessed_text)
        
        # Extract risk indicators for the whole document
        risk_indicators = self.extract_risk_indicators(
            preprocessed_text, 
            self.config["RISK_KEYWORDS"]
        )
        
        # Extract risk indicators for each section
        section_risk_indicators = {}
        for section_name, section_text in sections.items():
            section_risk_indicators[section_name] = self.extract_risk_indicators(
                section_text,
                self.config["RISK_KEYWORDS"]
            )
        
        # Create document features
        document_features = {
            "file_name": os.path.basename(file_path),
            "file_path": file_path,
            "file_size_mb": os.path.getsize(file_path) / (1024 * 1024),
            "document_length": len(raw_text),
            "sections": sections,
            "risk_indicators": risk_indicators,
            "section_risk_indicators": section_risk_indicators
        }
        
        return document_features